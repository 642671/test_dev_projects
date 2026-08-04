"""
Word文档模板与样式

负责生成格式化的Word文档，包含测试链路信息和截图
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class WordTemplates:
    """Word文档模板与样式"""
    
    @staticmethod
    def create_document():
        """
        创建新的Word文档并应用样式
        
        Returns:
            Document: python-docx Document实例
        """
        doc = Document()
        WordTemplates._apply_styles(doc)
        return doc
    
    @staticmethod
    def _apply_styles(doc):
        """
        应用全局样式
        
        Args:
            doc: Document实例
        """
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        font.color.rgb = RGBColor(51, 51, 51)
        
        # 设置中文字体
        rPr = style.element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        rPr.append(rFonts)
        
        # 设置段落间距
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.5
    
    @staticmethod
    def add_title(doc, title):
        """
        添加文档标题
        
        Args:
            doc: Document实例
            title: 标题文本
        """
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置标题样式
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.size = Pt(22)
    
    @staticmethod
    def add_chain_header(doc, chain_config):
        """
        添加链路头部信息（纯中文格式）
        
        Args:
            doc: Document实例
            chain_config: 链路配置字典
        """
        # 添加生成时间（右对齐）
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = subtitle.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # 空行
        
        # 创建基本信息表格（纯中文标签）
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 设置表格宽度
        for cell in table.columns[0].cells:
            cell.width = Cm(3)
        for cell in table.columns[1].cells:
            cell.width = Cm(13)
        
        # 填写基本信息（纯中文）
        meta_items = [
            ("链路名称", chain_config.get("chain_name", "")),
            ("环境地址", chain_config.get("environment", {}).get("url", "")),
            ("环境准备", chain_config.get("preparation", "无"))
        ]
        
        for i, (label, value) in enumerate(meta_items):
            WordTemplates._set_cell(table.cell(i, 0), label, bold=True, bg_color="D9E2F3")
            WordTemplates._set_cell(table.cell(i, 1), value)
        
        doc.add_paragraph()  # 空行
    
    @staticmethod
    def add_step(doc, step_result, step_index, screenshot_path=None):
        """
        添加单个步骤到文档（纯中文格式）
        
        Args:
            doc: Document实例
            step_result: 步骤执行结果字典
            step_index: 步骤序号
            screenshot_path: 截图文件路径（可选）
        """
        # 步骤标题（中文格式：第一步：xxx）
        step_name = step_result.get("name", f"步骤 {step_index}")
        
        # 转换数字为中文
        chinese_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if step_index <= 10:
            step_chinese = chinese_nums[step_index - 1]
        else:
            step_chinese = str(step_index)
        
        heading = doc.add_heading(f'第{step_chinese}步：{step_name}', level=2)
        
        # 设置标题颜色
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        # 步骤详情表格（纯中文标签）
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 设置表格宽度
        for cell in table.columns[0].cells:
            cell.width = Cm(3)
        for cell in table.columns[1].cells:
            cell.width = Cm(13)
        
        # 填写步骤详情（纯中文）
        details = [
            ("操作描述", step_result.get("description", "")),
            ("输入数据", step_result.get("input_data", "无")),
            ("预期结果", step_result.get("expected", ""))
        ]
        
        for i, (label, value) in enumerate(details):
            WordTemplates._set_cell(table.cell(i, 0), label, bold=True, bg_color="D9E2F3")
            WordTemplates._set_cell(table.cell(i, 1), value)
        
        # 添加截图占位符（留空，后续手动补充）
        doc.add_paragraph()  # 空行
        
        # 添加截图占位提示
        placeholder = doc.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = placeholder.add_run('【截图位置】')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.bold = True
        
        # 如果已有截图，也可以添加（可选）
        if screenshot_path and os.path.exists(screenshot_path):
            try:
                pic_paragraph = doc.add_paragraph()
                pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pic_paragraph.add_run()
                run.add_picture(screenshot_path, width=Inches(5.5))
            except Exception as e:
                pass  # 截图失败就只显示占位符
        
        # 添加分隔线
        doc.add_paragraph()
    
    @staticmethod
    def _set_cell(cell, text, bold=False, bg_color=None, font_color=None):
        """
        设置单元格内容和格式
        
        Args:
            cell: 单元格对象
            text: 文本内容
            bold: 是否粗体
            bg_color: 背景色（如 "D9E2F3"）
            font_color: 字体颜色（如 "FF0000"）
        """
        cell.text = ''
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        
        run.bold = bold
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        
        # 设置字体颜色
        if font_color:
            run.font.color.rgb = RGBColor(
                int(font_color[0:2], 16),
                int(font_color[2:4], 16),
                int(font_color[4:6], 16)
            )
        
        # 设置背景色
        if bg_color:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg_color)
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)
        
        # 设置段落间距
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        
        # 设置中文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        rPr.append(rFonts)
    
    @staticmethod
    def save_document(doc, output_path):
        """
        保存文档
        
        Args:
            doc: Document实例
            output_path: 输出文件路径
        """
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        
        doc.save(output_path)
        return output_path
